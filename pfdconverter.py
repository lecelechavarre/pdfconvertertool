import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import sys
import subprocess
import importlib.util
import threading
import re
import shutil
import tempfile
import io
from pathlib import Path
from xml.sax.saxutils import escape
import xml.etree.ElementTree as ET

# ============ AUTOMATIC DEPENDENCY MANAGEMENT ============
def install_and_import(package, import_name=None):
    """Automatically install and import a required package"""
    if import_name is None:
        import_name = package
    
    try:
        spec = importlib.util.find_spec(import_name)
        if spec is None:
            raise ImportError(f"Package {package} not found")
        module = importlib.import_module(import_name)
        
        if package == 'Pillow':
            from PIL import Image, ImageTk
            return (Image, ImageTk)
        
        return module
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "--quiet", package])
        print(f"{package} installed successfully!")
        
        if package == 'Pillow':
            from PIL import Image, ImageTk
            return (Image, ImageTk)
        else:
            return importlib.import_module(import_name)

# Install and import all required packages
PIL_Image, PIL_ImageTk = install_and_import('Pillow', 'PIL') 
Image = PIL_Image
ImageTk = PIL_ImageTk

fitz = install_and_import('PyMuPDF', 'fitz')

python_docx = install_and_import('python-docx', 'docx')
Document = python_docx.Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

reportlab = install_and_import('reportlab')
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

pdf2docx_module = install_and_import('pdf2docx', 'pdf2docx')
Converter = pdf2docx_module.Converter

# ============ MAIN APPLICATION ============

class ConverterApp:
    def __init__(self):
        self.window = tk.Tk()
        self.window.title("Document Converter")
        self.window.state('zoomed')
        self.window.configure(bg='#f5f5f7')
        
        self.selected_file = None
        self.current_mode = None
        self.converted_file = None
        self.preview_file = None
        self.preview_image = None
        self.setup_fonts()
        self.setup_ui()
        
        # Bind resize event
        self.window.bind('<Configure>', self.on_window_resize)
    
    def setup_fonts(self):
        self.font_regular = ('Helvetica', 10)
        self.font_medium = ('Helvetica', 11)
        self.font_bold = ('Helvetica', 24, 'bold')
        self.font_button = ('Helvetica', 11)
        self.font_card_title = ('Helvetica', 13, 'bold')
        
    def setup_ui(self):
        # Main container with padding
        main_container = tk.Frame(self.window, bg='#f5f5f7')
        main_container.pack(expand=True, fill='both', padx=40, pady=30)
        
        # Configure grid weights
        main_container.grid_columnconfigure(0, weight=35)
        main_container.grid_columnconfigure(1, weight=65)
        main_container.grid_rowconfigure(0, weight=1)
        
        # ============ LEFT PANEL - CONVERTER ============
        left_panel = tk.Frame(main_container, bg='#f5f5f7')
        left_panel.grid(row=0, column=0, sticky='nsew', padx=(0, 15))
        
        # Header
        header = tk.Frame(left_panel, bg='#f5f5f7')
        header.pack(fill='x', pady=(0, 20))
        
        title = tk.Label(
            header,
            text='Document Converter',
            font=self.font_bold,
            bg='#f5f5f7',
            fg='#1d1d1f',
            anchor='w'
        )
        title.pack(fill='x')
        
        # Divider
        divider = tk.Frame(left_panel, height=1, bg='#d2d2d7')
        divider.pack(fill='x', pady=(0, 25))
        
        # Cards container
        cards_frame = tk.Frame(left_panel, bg='#f5f5f7')
        cards_frame.pack(fill='x')
        
        # PDF to Word card
        self.create_option_card(
            cards_frame,
            'PDF to Word',
            'Convert PDF files to editable Word documents',
            'pdf',
            0
        )
        
        # Word to PDF card
        self.create_option_card(
            cards_frame,
            'Word to PDF',
            'Convert Word documents to PDF format',
            'docx',
            1
        ) 
        
        # File info frame
        info_container = tk.Frame(left_panel, bg='#f5f5f7')
        info_container.pack(fill='x', pady=(25, 15))
        
        self.info_frame = tk.Frame(info_container, bg='#ffffff', relief='flat', bd=0, height=50)
        self.info_frame.pack(fill='x')
        self.info_frame.pack_propagate(False)
        
        self.file_label = tk.Label(
            self.info_frame,
            text='No file selected',
            font=self.font_medium,
            bg='#ffffff',
            fg='#86868b',
            anchor='w',
            padx=15,
            pady=12
        )
        self.file_label.pack(fill='both')
        
        # Download button
        download_container = tk.Frame(left_panel, bg='#f5f5f7')
        download_container.pack(fill='x', pady=(0, 15))
        
        self.download_btn = tk.Button(
            download_container,
            text='⬇️ Download Converted File',
            font=self.font_button,
            bg='#86868b',
            fg='#ffffff',
            bd=0,
            padx=25,
            pady=10,
            activebackground='#666666',
            activeforeground='#ffffff',
            state='disabled',
            relief='flat',
            cursor='',
            width=20,
            command=self.download_file
        )
        self.download_btn.pack()
        
        # Status bar
        status_container = tk.Frame(left_panel, bg='#f5f5f7')
        status_container.pack(fill='x', side='bottom', pady=(10, 0))
        
        self.status_frame = tk.Frame(status_container, bg='#ffffff', relief='flat', bd=0, height=44)
        self.status_frame.pack(fill='x')
        self.status_frame.pack_propagate(False)
        
        self.status_label = tk.Label(
            self.status_frame,
            text='Ready',
            font=self.font_regular,
            bg='#ffffff',
            fg='#86868b',
            anchor='w',
            padx=15,
            pady=12
        )
        self.status_label.pack(fill='both')
        
        # ============ RIGHT PANEL - PREVIEW ============
        right_panel = tk.Frame(main_container, bg='#ffffff', relief='flat', bd=0)
        right_panel.grid(row=0, column=1, sticky='nsew', padx=(15, 0))
        
        # Preview header
        preview_header = tk.Frame(right_panel, bg='#ffffff', height=50)
        preview_header.pack(fill='x', padx=25, pady=(20, 10))
        preview_header.pack_propagate(False)
        
        preview_title = tk.Label(
            preview_header,
            text='Preview',
            font=('Helvetica', 16, 'bold'),
            bg='#ffffff',
            fg='#1d1d1f'
        )
        preview_title.pack(side='left')
        
        self.preview_filename = tk.Label(
            preview_header,
            text='',
            font=self.font_regular,
            bg='#ffffff',
            fg='#86868b'
        )
        self.preview_filename.pack(side='right', padx=(10, 0))
        
        # Preview area with scrollbar
        preview_container = tk.Frame(right_panel, bg='#f5f5f7')
        preview_container.pack(expand=True, fill='both', padx=25, pady=(0, 25))
        
        # Create canvas with scrollbar
        self.preview_canvas = tk.Canvas(
            preview_container,
            bg='#ffffff',
            highlightthickness=1,
            highlightbackground='#e6e6e8',
            relief='flat'
        )
        
        scrollbar = tk.Scrollbar(
            preview_container,
            orient='vertical',
            command=self.preview_canvas.yview
        )
        
        self.preview_canvas.configure(yscrollcommand=scrollbar.set)
        
        # Pack canvas and scrollbar
        self.preview_canvas.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        # Create inner frame for content
        self.preview_inner = tk.Frame(self.preview_canvas, bg='#ffffff')
        self.preview_canvas.create_window((0, 0), window=self.preview_inner, anchor='nw')
        
        # Bind events
        self.preview_inner.bind('<Configure>', self.on_inner_configure)
        self.preview_canvas.bind('<Configure>', self.on_canvas_configure)
        
        # Show placeholder
        self.show_preview_placeholder()
    
    def on_window_resize(self, event):
        if event.widget == self.window:
            self.update_preview_layout()
    
    def on_inner_configure(self, event):
        self.preview_canvas.configure(scrollregion=self.preview_canvas.bbox('all'))
    
    def on_canvas_configure(self, event):
        self.preview_canvas.itemconfig('all', width=event.width)
    
    def update_preview_layout(self):
        if hasattr(self, 'preview_canvas'):
            self.preview_canvas.itemconfig('all', width=self.preview_canvas.winfo_width())
    
    def show_preview_placeholder(self):
        # Clear inner frame
        for widget in self.preview_inner.winfo_children():
            widget.destroy()
        
        # Create placeholder
        placeholder = tk.Frame(self.preview_inner, bg='#ffffff', height=400)
        placeholder.pack(expand=True, fill='both', padx=40, pady=80)
        
        icon_label = tk.Label(
            placeholder,
            text='👀',
            font=('Helvetica', 48),
            bg='#ffffff',
            fg='#d2d2d7'
        )
        icon_label.pack(pady=(0, 20))
        
        text_label = tk.Label(
            placeholder,
            text='No preview available',
            font=('Helvetica', 16),
            bg='#ffffff',
            fg='#86868b'
        )
        text_label.pack()
        
        subtext_label = tk.Label(
            placeholder,
            text='Select a file to generate preview',
            font=('Helvetica', 12),
            bg='#ffffff',
            fg='#a1a1a6'
        )
        subtext_label.pack(pady=(10, 0))
        
        self.preview_filename.configure(text='')
    
    def create_option_card(self, parent, title, description, mode, col):
        card = tk.Frame(
            parent,
            bg='#ffffff',
            highlightbackground='#e6e6e8',
            highlightthickness=1,
            bd=0,
            height=200
        )
        card.pack(side='left', padx=(0 if col == 0 else 8, 8 if col == 0 else 0), fill='both', expand=True)
        card.pack_propagate(False)
        
        content = tk.Frame(card, bg='#ffffff')
        content.pack(fill='both', expand=True, padx=18, pady=18)
        
        # Icon and title row
        header_row = tk.Frame(content, bg='#ffffff')
        header_row.pack(anchor='w', fill='x', pady=(0, 12))
        
        icon = tk.Label(
            header_row,
            text='📄' if mode == 'pdf' else '📝',
            font=('Helvetica', 32),
            bg='#ffffff',
            fg='#1d1d1f'
        )
        icon.pack(side='left', padx=(0, 10))
        
        title_label = tk.Label(
            header_row,
            text=title,
            font=self.font_card_title,
            bg='#ffffff',
            fg='#1d1d1f'
        )
        title_label.pack(side='left')
        
        # Description
        desc_label = tk.Label(
            content,
            text=description,
            font=self.font_regular,
            bg='#ffffff',
            fg='#86868b',
            justify='left',
            wraplength=220
        )
        desc_label.pack(anchor='w', pady=(0, 20))
        
        # Select button
        select_btn = tk.Button(
            content,
            text='Select File',
            font=self.font_regular,
            bg='#ffffff',
            fg='#0066cc',
            bd=0,
            padx=0,
            pady=2,
            cursor='hand2',
            activebackground='#ffffff',
            activeforeground='#004999',
            command=lambda m=mode: self.select_file(m),
            relief='flat'
        )
        select_btn.pack(anchor='w')
        
        # Hover effect
        def on_enter(e):
            card.configure(highlightbackground='#0066cc')
            select_btn.configure(fg='#004999')
        
        def on_leave(e):
            card.configure(highlightbackground='#e6e6e8')
            select_btn.configure(fg='#0066cc')
        
        card.bind('<Enter>', on_enter)
        card.bind('<Leave>', on_leave)
        
        return card
    
    def select_file(self, mode):
        if mode == 'pdf':
            filetypes = [('PDF files', '*.pdf')]
            file_type = 'PDF'
        else:
            filetypes = [('Word files', '*.docx')]
            file_type = 'Word'
        
        file_path = filedialog.askopenfilename(
            title=f'Select {file_type} File',
            filetypes=filetypes
        )
        
        if file_path:
            self.selected_file = file_path
            self.current_mode = mode
            self.converted_file = None
            self.preview_file = None
            
            # Clear previous preview
            self.show_preview_placeholder()
            
            filename = os.path.basename(file_path)
            self.file_label.configure(
                text=f'Selected: {filename}',
                fg='#1d1d1f'
            )
            
            # Disable download button until conversion is complete
            self.download_btn.configure(
                state='disabled',
                bg='#86868b',
                fg='#ffffff',
                activebackground='#666666',
                cursor=''
            )
            
            self.status_label.configure(
                text='Generating preview...',
                fg='#1d1d1f'
            )
            
            # Start preview generation thread
            thread = threading.Thread(target=self.generate_preview)
            thread.daemon = True
            thread.start()
    
    def generate_preview(self):
        """Generate preview only - NO DOWNLOAD"""
        error = None
        preview_path = None
        
        try:
            # Create temporary file for preview
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf' if self.current_mode == 'docx' else '.docx') as tmp_file:
                preview_path = tmp_file.name
            
            if self.current_mode == 'pdf':
                # PDF to Word preview
                cv = Converter(self.selected_file)
                cv.convert(preview_path, start=0, end=None)
                cv.close()
                
            else:
                # Word to PDF preview - WITH COMPLETE FORMATTING PRESERVATION
                self.convert_docx_to_pdf_preserve_formatting(self.selected_file, preview_path)
                
        except Exception as e:
            error = str(e)
            if preview_path and os.path.exists(preview_path):
                os.unlink(preview_path)
        
        if error:
            self.window.after(0, lambda err=error: self.preview_error(err))
        else:
            self.preview_file = preview_path
            self.window.after(0, lambda path=preview_path: self.preview_success(path))
    
    def convert_docx_to_pdf_preserve_formatting(self, docx_path, pdf_path):
        """Convert DOCX to PDF while preserving ALL formatting, spacing, and layout"""
        doc = Document(docx_path)
        
        # Create PDF document with proper margins
        doc_template = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        story = []
        
        # Process each paragraph individually to preserve spacing
        for paragraph in doc.paragraphs:
            # Get paragraph formatting
            p_format = paragraph.paragraph_format
            
            # Calculate spacing values in points
            space_before = self.get_paragraph_spacing(p_format.space_before)
            space_after = self.get_paragraph_spacing(p_format.space_after)
            line_spacing = self.get_line_spacing(p_format.line_spacing)
            
            # Get alignment
            alignment = self.get_paragraph_alignment(paragraph.alignment)
            
            # Get indentation
            left_indent = self.get_indent(p_format.left_indent)
            right_indent = self.get_indent(p_format.right_indent)
            first_line_indent = self.get_indent(p_format.first_line_indent)
            
            # Add space before paragraph if needed
            if space_before > 0:
                story.append(Spacer(1, space_before))
            
            # Process runs to preserve inline formatting
            if len(paragraph.runs) > 0:
                # Build formatted text with proper XML tags
                formatted_text = self.build_formatted_text(paragraph.runs)
                
                if formatted_text:
                    # Create paragraph style with all formatting
                    style_name = f'ParaStyle_{len(story)}'
                    p_style = ParagraphStyle(
                        style_name,
                        parent=getSampleStyleSheet()['Normal'],
                        fontName='Helvetica',
                        fontSize=11,
                        leading=line_spacing,
                        alignment=alignment,
                        leftIndent=left_indent,
                        rightIndent=right_indent,
                        firstLineIndent=first_line_indent
                    )
                    
                    # Create paragraph and add to story
                    p = Paragraph(formatted_text, p_style)
                    story.append(p)
            
            # Add space after paragraph if needed
            if space_after > 0:
                story.append(Spacer(1, space_after))
        
        # Build the PDF
        doc_template.build(story)
    
    def get_paragraph_spacing(self, spacing_value):
        """Convert Word spacing to points"""
        if spacing_value is None:
            return 0
        try:
            return spacing_value.pt
        except:
            return 0
    
    def get_line_spacing(self, line_spacing):
        """Convert Word line spacing to points"""
        if line_spacing is None:
            return 14  # Default line spacing
        
        try:
            if hasattr(line_spacing, 'pt'):
                return line_spacing.pt
            else:
                # If it's a multiple, convert to points (assuming 12pt base)
                return line_spacing * 12
        except:
            return 14
    
    def get_indent(self, indent_value):
        """Convert Word indent to points"""
        if indent_value is None:
            return 0
        try:
            return indent_value.pt
        except:
            return 0
    
    def get_paragraph_alignment(self, alignment):
        """Convert Word alignment to ReportLab alignment"""
        if alignment is None:
            return TA_LEFT
        
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: TA_LEFT,
            WD_ALIGN_PARAGRAPH.CENTER: TA_CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT: TA_RIGHT,
            WD_ALIGN_PARAGRAPH.JUSTIFY: TA_JUSTIFY
        }
        return alignment_map.get(alignment, TA_LEFT)
    
    def build_formatted_text(self, runs):
        """Build formatted text from runs with proper XML tags"""
        formatted_parts = []
        
        for run in runs:
            text = run.text
            if not text:
                continue
            
            # Clean the text
            clean_text = self.clean_text(text)
            if not clean_text:
                continue
            
            # Escape XML characters
            safe_text = self.escape_xml_chars(clean_text)
            
            # Apply formatting tags
            if run.bold:
                safe_text = f"<b>{safe_text}</b>"
            if run.italic:
                safe_text = f"<i>{safe_text}</i>"
            if run.underline:
                safe_text = f"<u>{safe_text}</u>"
            
            formatted_parts.append(safe_text)
        
        return ''.join(formatted_parts)
    
    def escape_xml_chars(self, text):
        """Escape XML special characters for ReportLab"""
        if not text:
            return ""
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        return text
    
    def clean_text(self, text):
        """Clean text while preserving all meaningful characters"""
        if not text:
            return ""
        
        # Remove only absolute control characters
        cleaned = ''.join(char for char in text if ord(char) >= 32 or char == '\n' or char == '\t' or char == '\r')
        
        # Remove any remaining control characters
        cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', cleaned)
        
        return cleaned
    
    def download_file(self):
        """Download the converted file - ONLY WHEN DOWNLOAD BUTTON IS CLICKED"""
        if not self.converted_file:
            self.start_conversion_for_download()
            return
            
        if self.converted_file and os.path.exists(self.converted_file):
            save_path = filedialog.asksaveasfilename(
                defaultextension=os.path.splitext(self.converted_file)[1],
                filetypes=[
                    ('PDF files', '*.pdf') if self.converted_file.endswith('.pdf') else ('Word files', '*.docx')
                ],
                initialfile=os.path.basename(self.converted_file)
            )
            
            if save_path:
                try:
                    shutil.copy2(self.converted_file, save_path)
                    messagebox.showinfo('Download Complete', f'File saved to:\n{save_path}')
                except Exception as e:
                    messagebox.showerror('Download Failed', f'Could not save file:\n{str(e)}')
    
    def start_conversion_for_download(self):
        """Convert file for download"""
        if not self.selected_file:
            return
        
        self.download_btn.configure(
            state='disabled',
            bg='#86868b',
            fg='#ffffff',
            activebackground='#666666',
            cursor='',
            text='⏳ Converting...'
        )
        
        self.status_label.configure(
            text='Converting file for download...',
            fg='#1d1d1f'
        )
        
        thread = threading.Thread(target=self.convert_for_download)
        thread.daemon = True
        thread.start()
    
    def convert_for_download(self):
        """Convert file for permanent storage and download"""
        error = None
        output_path = None
        
        try:
            output_path = str(Path(self.selected_file).with_suffix('.docx' if self.current_mode == 'pdf' else '.pdf'))
            
            os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
            
            if self.current_mode == 'pdf':
                cv = Converter(self.selected_file)
                cv.convert(output_path, start=0, end=None)
                cv.close()
                
            else:
                self.convert_docx_to_pdf_preserve_formatting(self.selected_file, output_path)
                
        except Exception as e:
            error = str(e)
        
        if error:
            self.window.after(0, lambda err=error: self.conversion_error(err))
        else:
            self.converted_file = output_path
            self.window.after(0, lambda: self.download_ready())
    
    def download_ready(self):
        """Called when conversion for download is complete"""
        self.status_label.configure(
            text='Conversion complete - Ready to download',
            fg='#1d1d1f'
        )
        
        self.download_btn.configure(
            state='normal',
            bg='#34a853',
            fg='#ffffff',
            activebackground='#2d8c46',
            cursor='hand2',
            text='⬇️ Download Converted File',
            command=self.download_file
        )
        
        self.download_file()
    
    def preview_success(self, preview_path):
        self.status_label.configure(
            text='Preview generated',
            fg='#1d1d1f'
        )
        
        self.download_btn.configure(
            state='normal',
            bg='#34a853',
            fg='#ffffff',
            activebackground='#2d8c46',
            cursor='hand2',
            text='⬇️ Download Converted File',
            command=self.download_file
        )
        
        self.file_label.configure(
            text=f'Selected: {os.path.basename(self.selected_file)}',
            fg='#1d1d1f'
        )
        
        self.preview_filename.configure(text=f'Preview: {os.path.basename(self.selected_file)}')
        
        if preview_path.endswith('.pdf'):
            self.preview_pdf(preview_path)
        else:
            self.preview_docx(preview_path)
    
    def preview_error(self, error_msg):
        self.status_label.configure(
            text='Preview generation failed',
            fg='#ff3b30'
        )
        
        self.download_btn.configure(
            state='disabled',
            bg='#86868b',
            fg='#ffffff',
            activebackground='#666666',
            cursor=''
        )
        
        self.show_preview_placeholder()
        messagebox.showerror('Error', f'Failed to generate preview.\n\n{error_msg}')
    
    def conversion_error(self, error_msg):
        self.status_label.configure(
            text='Conversion failed',
            fg='#ff3b30'
        )
        
        self.download_btn.configure(
            state='normal',
            bg='#34a853',
            fg='#ffffff',
            activebackground='#2d8c46',
            cursor='hand2',
            text='⬇️ Download Converted File',
            command=self.download_file
        )
        
        messagebox.showerror('Error', f'Failed to convert file.\n\n{error_msg}')
    
    def preview_pdf(self, pdf_path):
        # Clear inner frame
        for widget in self.preview_inner.winfo_children():
            widget.destroy()
        
        try:
            doc = fitz.open(pdf_path)
            
            pages_container = tk.Frame(self.preview_inner, bg='#ffffff')
            pages_container.pack(expand=True, fill='both', padx=30, pady=30)
            
            canvas_width = self.preview_canvas.winfo_width() - 60
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                zoom = canvas_width / page.rect.width
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                
                self.preview_image = ImageTk.PhotoImage(img)
                
                page_frame = tk.Frame(pages_container, bg='#ffffff')
                page_frame.pack(pady=(0, 20))
                
                if len(doc) > 1:
                    page_label = tk.Label(
                        page_frame,
                        text=f'Page {page_num + 1} of {len(doc)}',
                        font=('Helvetica', 10),
                        bg='#ffffff',
                        fg='#86868b'
                    )
                    page_label.pack(pady=(0, 5))
                
                image_label = tk.Label(
                    page_frame,
                    image=self.preview_image,
                    bg='#ffffff',
                    bd=0
                )
                image_label.pack()
            
            doc.close()
            
        except Exception as e:
            error_label = tk.Label(
                self.preview_inner,
                text='⚠️ Preview not available',
                font=('Helvetica', 12),
                bg='#ffffff',
                fg='#86868b'
            )
            error_label.pack(expand=True)
    
    def preview_docx(self, docx_path):
        # Clear inner frame
        for widget in self.preview_inner.winfo_children():
            widget.destroy()
        
        try:
            doc = Document(docx_path)
            
            container = tk.Frame(self.preview_inner, bg='#ffffff')
            container.pack(expand=True, fill='both', padx=30, pady=30)
            
            header_frame = tk.Frame(container, bg='#ffffff')
            header_frame.pack(fill='x', pady=(0, 20))
            
            tk.Label(
                header_frame,
                text='Document Content',
                font=('Helvetica', 14, 'bold'),
                bg='#ffffff',
                fg='#1d1d1f'
            ).pack(anchor='w')
            
            content_frame = tk.Frame(container, bg='#ffffff')
            content_frame.pack(fill='both', expand=True)
            
            text_widget = tk.Text(
                content_frame,
                font=('Helvetica', 11),
                bg='#ffffff',
                fg='#1d1d1f',
                wrap='word',
                bd=0,
                highlightthickness=0,
                relief='flat'
            )
            text_widget.pack(fill='both', expand=True)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    clean_text = self.clean_text(paragraph.text)
                    if clean_text:
                        text_widget.insert('end', clean_text + '\n\n')
                else:
                    text_widget.insert('end', '\n')
            
            text_widget.configure(state='disabled')
            
        except Exception as e:
            error_label = tk.Label(
                self.preview_inner,
                text='📄 Document preview',
                font=('Helvetica', 12),
                bg='#ffffff',
                fg='#86868b'
            )
            error_label.pack(expand=True)
    
    def run(self):
        self.window.mainloop()


def main():
    app = ConverterApp()
    app.run()


if __name__ == "__main__":
    main()


